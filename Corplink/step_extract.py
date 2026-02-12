# coding: utf-8
import os
import re
from pathlib import Path
from typing import List, Dict

from docx import Document
from tqdm import tqdm
import numpy as np
from .constants import BASE_DIR
import pandas as pd

from . import state
from .factiva_rtf import read_rtf_text, parse_records_from_text
from .options import ExtractMode

from .constants import DATE_FINDER, ANCHOR_TEXT, BASE_DIR
from .env_bootstrap import cute_box
from .text_utils import _normalize, clean_text
from . import state
from .model_utils import model_emb

def extract_sentences(path: Path) -> List[str]:
    doc = Document(path)
    collecting, current, articles = False, "", []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt: 
            continue
        tag = txt.lower()
        if tag == "body": 
            collecting, current = True, ""
            continue
        if tag in ("notes", "classification") and collecting:
            collecting = False
            articles.append(current.strip())
            continue
        if collecting: 
            current += " " + txt
    sents = []
    for art in articles:
        for s in re.split(r"\.\s*", art):
            s = s.strip()
            if len(s) >= 5: 
                sents.append(s)
    return sents

def extract_index_titles(paragraphs):
    paras_text = [p.text.strip() for p in paragraphs]
    m = re.search(r'Documents?\s*\(\s*(\d+)\s*\)', '\n'.join(paras_text[:50]), re.I)
    if not m: 
        return []
    total = int(m.group(1))
    
    pat = re.compile(r'^(\d+)\.\s+(.*)$')
    seen = set()
    titles = []
    
    for i, line in enumerate(paras_text):
        m2 = pat.match(line)
        if m2:
            is_valid_toc = False
            for offset in range(1, 7): 
                if i + offset >= len(paras_text): break
                next_line = paras_text[i + offset].lower()
                if "client/matter" in next_line or "search terms" in next_line:
                    is_valid_toc = True
                    break
            if not is_valid_toc:
                continue 
            
            raw = m2.group(2).strip()
            norm = _normalize(raw)
            
            if norm in seen: 
                continue
            seen.add(norm)
            titles.append((int(m2.group(1)), raw, norm))
            
            if len(titles) >= total: 
                break
                
    return sorted(titles, key=lambda x: x[0])

def _filter_sentences(raw_sents: List[str]) -> List[Dict]:
    recs: List[Dict] = []
    if state.USE_SEMANTIC_FILTER and raw_sents:
        if not hasattr(_filter_sentences, "anchor_vec"):
            _filter_sentences.anchor_vec = model_emb.encode([ANCHOR_TEXT], normalize_embeddings=True)[0]
        sent_vecs = model_emb.encode(raw_sents, normalize_embeddings=True)
        sim_scores = np.dot(sent_vecs, _filter_sentences.anchor_vec)
    else:
        sim_scores = [0.0] * len(raw_sents)

    for i, sent in enumerate(raw_sents):
        is_hit = False
        match_reason = ""
        hit_count = 0

        if state.USE_SEMANTIC_FILTER:
            score = float(sim_scores[i])
            if score > 0.45:
                is_hit = True
                hit_count = 1
                match_reason = f"Semantic({score:.2f})"
        else:
            hits = [k for k in state.KEYWORD_ROOTS if k in sent.lower()]
            if hits:
                is_hit = True
                hit_count = len(hits)
                match_reason = "; ".join(hits)

        if is_hit:
            recs.append({
                "Sentence": sent,
                "Hit_Count": hit_count,
                "Matched_Keywords": match_reason
            })
    return recs

def extract_sentences_by_titles(filepath: str) -> List[Dict]:
    doc = Document(filepath); paras = doc.paragraphs
    index_titles = extract_index_titles(paras); recs = []
    
    if index_titles:
        paras_norm = [_normalize(p.text) for p in paras]
        last_article_end_idx = 0

        for i_title, (doc_idx, title_raw, title_norm) in enumerate(index_titles):
            match_idx = -1
            date_line_idx = -1
            
            candidates = [i for i, n in enumerate(paras_norm) 
                          if i >= last_article_end_idx and n == title_norm]
            
            for idx in candidates:
                if idx + 1 < len(paras):
                    next_line = paras[idx+1].text.strip().lower()
                    if next_line.startswith("client/matter") or next_line.startswith("search terms"):
                        continue 

                found_date = False
                temp_date_idx = -1
                for offset in range(1, 4):
                    if idx + offset >= len(paras): break
                    txt = paras[idx + offset].text.strip()
                    if DATE_FINDER.search(txt):
                        found_date = True
                        temp_date_idx = idx + offset
                        break
                
                if found_date:
                    match_idx = idx
                    date_line_idx = temp_date_idx
                    break
            
            if match_idx == -1: 
                continue

            if date_line_idx > match_idx + 1:
                publisher = paras[match_idx + 1].text.strip()
            else:
                publisher = ""
            
            news_date = ""
            m = DATE_FINDER.search(paras[date_line_idx].text.strip())
            if m: 
                news_date = m.group(0)

            pub_idx = date_line_idx 
            search_end_limit = len(paras)
            if i_title + 1 < len(index_titles):
                next_title_norm = index_titles[i_title+1][2]
                try:
                    next_candidates = [i for i, n in enumerate(paras_norm) 
                                       if i > match_idx + 20 and n == next_title_norm]
                    if next_candidates: 
                        search_end_limit = next_candidates[0]
                except Exception: 
                    pass

            body_start = next((i+1 for i in range(pub_idx+1, search_end_limit) if paras[i].text.strip().lower() == "body"), None)
            if body_start is None: 
                body_start = pub_idx + 1
            
            body_end = len(paras)
            for i in range(body_start, search_end_limit):
                t_low = paras[i].text.strip().lower()
                if t_low.startswith("notes") or t_low.startswith("classification") or "(end) dow jones" in t_low:
                    body_end = i
                    break
            last_article_end_idx = body_end

            article = " ".join(clean_text(paras[i].text) for i in range(body_start, body_end))
            raw_sents = [s.strip() for s in re.split(r"\.\s*", article) if len(s.strip())>=20]

            if state.USE_SEMANTIC_FILTER and raw_sents:
                if not hasattr(extract_sentences_by_titles, "anchor_vec"):
                     extract_sentences_by_titles.anchor_vec = model_emb.encode([ANCHOR_TEXT], normalize_embeddings=True)[0]
                
                sent_vecs = model_emb.encode(raw_sents, normalize_embeddings=True)
                sim_scores = np.dot(sent_vecs, extract_sentences_by_titles.anchor_vec)
            else:
                sim_scores = [0.0] * len(raw_sents)

            for i, sent in enumerate(raw_sents):
                is_hit = False
                match_reason = ""
                hit_count = 0

                if state.USE_SEMANTIC_FILTER:
                    score = float(sim_scores[i])
                    if score > 0.45:
                        is_hit = True
                        hit_count = 1
                        match_reason = f"Semantic({score:.2f})"
                else:
                    hits = [k for k in state.KEYWORD_ROOTS if k in sent.lower()]
                    if hits:
                        is_hit = True
                        hit_count = len(hits)
                        match_reason = "; ".join(hits)
                
                if is_hit:
                    recs.append({
                        "Title": title_raw,
                        "Publisher": publisher,
                        "Date": news_date,
                        "Country": "",
                        "Sentence": sent,
                        "Hit_Count": hit_count,
                        "Matched_Keywords": match_reason
                    })
        
        if recs: 
            return recs

    for sent in extract_sentences(Path(filepath)):
        hits = [k for k in state.KEYWORD_ROOTS if k in sent.lower()]
        if hits:
             recs.append({
                "Title": "", "Publisher": "", "Date": "", "Country": "", 
                "Sentence": sent, "Hit_Count": len(hits), 
                "Matched_Keywords": "; ".join(hits)
            })
    return recs

def extract_sentences_from_factiva(filepath: str) -> List[Dict]:
    text = read_rtf_text(Path(filepath))
    records = parse_records_from_text(text)
    recs: List[Dict] = []

    for record in records:
        # å¥å­åˆ‡åˆ†é€»è¾‘ä¸ Lexis ä¿æŒä¸€è‡´
        raw_sents = [s.strip() for s in re.split(r"\.\s*", record.body) if len(s.strip()) >= 20]

        for hit in _filter_sentences(raw_sents):
            recs.append({
                "Title": record.title,
                "Publisher": record.publisher,
                "Date": record.date_yyyy_mm_dd,
                "Country": "",
                "Sentence": hit["Sentence"],
                "Hit_Count": hit["Hit_Count"],
                "Matched_Keywords": hit["Matched_Keywords"]
            })
    return recs
    
def step1():
    cute_box(
        "Step-1ï¼šæå– Word å¥å­ ä¸­â€¦",
        "Step-1ï¼šæ–‡æŠ½å‡ºä¸­â€¦",
        "ğŸ“„"
    )
    all_recs: List[Dict] = []

    if state.EXTRACT_MODE == ExtractMode.FACTIVA.value:
        rtf_files = []
        for root, _, files in os.walk(BASE_DIR):
            for fname in files:
                if not fname.endswith(".rtf") or fname.startswith("~$"):
                    continue
                full = Path(root) / fname
                rel = full.relative_to(BASE_DIR).parts
                tier1 = rel[0] if len(rel) >= 1 else ""
                tier2 = rel[1] if len(rel) >= 2 else ""
                rtf_files.append((str(full), tier1, tier2, fname))

        for fp, t1, t2, fname in tqdm(rtf_files, desc="ğŸ—‚ï¸ å¤„ç† Factiva RTF æ–‡ä»¶"):
            for r in extract_sentences_from_factiva(fp):
                r.update({"Tier_1": t1, "Tier_2": t2, "Filename": fname})
                all_recs.append(r)
    else:
        docx_files = []
        for root, _, files in os.walk(BASE_DIR):
            for fname in files:
                if not fname.endswith(".docx") or fname.startswith("~$"):
                    continue
                full = Path(root) / fname
                rel = full.relative_to(BASE_DIR).parts
                tier1 = rel[0] if len(rel) >= 1 else ""
                tier2 = rel[1] if len(rel) >= 2 else ""
                docx_files.append((str(full), tier1, tier2, fname))

        for fp, t1, t2, fname in tqdm(docx_files, desc="ğŸ“„ å¤„ç† Word æ–‡ä»¶"):
            for r in extract_sentences_by_titles(fp):
                if not r["Title"]:
                    r["Title"] = Path(fname).stem
                r.update({"Tier_1": t1, "Tier_2": t2, "Filename": fname})
                all_recs.append(r)

    state.SENTENCE_RECORDS = all_recs
    cute_box(
        f"Step-1 å®Œæˆï¼Œå…± {len(all_recs)} æ¡è®°å½•",
        f"Step-1 å®Œäº†ã—ã¾ã—ãŸï¼šå…¨{len(all_recs)}ä»¶",
        "âœ…"
    )
